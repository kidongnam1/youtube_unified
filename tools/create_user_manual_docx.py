from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
OUT_PATH = OUT_DIR / "YouTube_Unified_사용_매뉴얼.docx"
TXT_PATH = OUT_DIR / "YouTube_Unified_사용_매뉴얼_UTF8.txt"


TEXT_LINES: list[str] = []
STEP_COUNTER = 0


def record(text: str = "") -> None:
    TEXT_LINES.append(text)


def set_korean_font(run, size: int | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    record(title)
    record(subtitle)
    record()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    set_korean_font(run, 24)
    run.font.color.rgb = RGBColor(20, 38, 70)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_korean_font(run, 11)
    run.font.color.rgb = RGBColor(90, 100, 120)


def add_h1(doc: Document, text: str) -> None:
    global STEP_COUNTER
    STEP_COUNTER = 0
    record()
    record(text)
    record("=" * len(text))
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        set_korean_font(run, 16)


def add_h2(doc: Document, text: str) -> None:
    global STEP_COUNTER
    STEP_COUNTER = 0
    record()
    record(text)
    record("-" * len(text))
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_korean_font(run, 13)


def add_body(doc: Document, text: str) -> None:
    record(text)
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        set_korean_font(run)


def add_step(doc: Document, text: str) -> None:
    global STEP_COUNTER
    STEP_COUNTER += 1
    record(f"{STEP_COUNTER}. {text}")
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_korean_font(run)


def add_bullet(doc: Document, text: str) -> None:
    record(f"- {text}")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        set_korean_font(run)


def add_note(doc: Document, text: str) -> None:
    record(f"참고: {text}")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run("참고: ")
    run.bold = True
    set_korean_font(run)
    run.font.color.rgb = RGBColor(180, 90, 0)
    run = p.add_run(text)
    set_korean_font(run)


def add_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    record()
    record("항목 | 설명")
    record("--- | ---")
    for name, desc in rows:
        record(f"{name} | {desc}")
    record()
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "항목"
    hdr[1].text = "설명"
    for name, desc in rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = desc
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_korean_font(run)
    doc.add_paragraph()


def build_manual() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    TEXT_LINES.clear()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Malgun Gothic"
    styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Heading 2"].font.name = "Malgun Gothic"
    styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    add_title(
        doc,
        "YouTube Unified 사용 매뉴얼",
        "프로그램 실행부터 자막 추출, 스토리보드 제작, 이미지 생성, 영상 내보내기까지",
    )
    add_body(
        doc,
        "이 문서는 오랜만에 프로그램을 다시 열어도 순서대로 따라 할 수 있도록 단계별로 작성한 매뉴얼입니다.",
    )

    add_h1(doc, "1. 프로그램을 처음 여는 방법")
    add_step(doc, "폴더 D:\\program_kdn\\youtube_unified 로 이동합니다.")
    add_step(doc, "RUN_MENU.bat 파일을 더블클릭합니다. PowerShell을 선호하면 RUN_MENU.ps1을 실행합니다.")
    add_step(doc, "작은 런처 창이 뜨고, 통합 웹 앱이 자동으로 시작됩니다.")
    add_step(doc, "브라우저가 열리면 주소가 http://127.0.0.1:3000/ 인지 확인합니다.")
    add_step(doc, "화면 왼쪽에 YouTube Unified 메뉴가 보이면 정상입니다.")
    add_note(doc, "브라우저가 자동으로 열리지 않으면 직접 http://127.0.0.1:3000/ 주소를 입력합니다.")

    add_h1(doc, "2. 가장 먼저 해야 하는 API 설정")
    add_body(doc, "스토리보드 생성과 이미지 생성은 Gemini API Key가 있어야 동작합니다.")
    add_step(doc, "왼쪽 메뉴 아래의 API 설정 영역을 찾습니다.")
    add_step(doc, "Gemini API Key 입력칸에 키를 붙여넣습니다.")
    add_step(doc, "텍스트 모델은 기본값 gemini-2.5-flash를 그대로 둡니다.")
    add_step(doc, "이미지 모델은 gemini-2.5-flash-image를 그대로 둡니다.")
    add_step(doc, "이미지 화풍을 고릅니다. 처음에는 한국 경제 카툰 또는 화풍 없음을 권장합니다.")
    add_step(doc, "저장 버튼을 누릅니다.")
    add_step(doc, "테스트 버튼을 눌러 키 정상 메시지가 나오는지 확인합니다.")
    add_table(
        doc,
        [
            ("Gemini API Key", "스토리보드 생성, 이미지 생성, 일부 텍스트 처리에 필요합니다."),
            ("텍스트 모델", "대본을 분석하고 씬을 만드는 모델입니다."),
            ("이미지 모델", "각 씬의 대표 이미지를 만드는 모델입니다."),
            ("이미지 화풍", "생성 이미지의 전체 스타일을 정합니다."),
            ("커스텀 화풍", "직접 만든 영어 스타일 프롬프트를 적용합니다."),
        ],
    )

    add_h1(doc, "3. 작업 목적별 빠른 길")
    add_body(doc, "원하는 작업에 따라 아래 중 하나를 선택해서 진행합니다.")
    add_table(
        doc,
        [
            ("새 영상 기획", "스토리보드 생성 탭에서 주제를 입력하고 스토리보드를 만듭니다."),
            ("유튜브 자막 활용", "자막 추출 탭에서 자막을 가져온 뒤 스토리보드로 보냅니다."),
            ("기존 영상 분석", "영상 자동화 탭에서 URL 또는 파일 경로를 넣고 자동화를 시작합니다."),
            ("이전 작업 이어하기", "저장된 프로젝트 탭에서 불러오기를 누릅니다."),
            ("결과 파일 저장", "스토리보드 생성 탭에서 JSON, CSV, MP4/WebM으로 내보냅니다."),
        ],
    )

    add_h1(doc, "4. 새 주제로 스토리보드 만드는 절차")
    add_step(doc, "왼쪽 메뉴에서 스토리보드 생성 탭을 누릅니다.")
    add_step(doc, "작업 이름에 알아보기 쉬운 이름을 입력합니다. 예: 2026 경제 이슈 쇼츠")
    add_step(doc, "주제에 만들 영상 주제를 입력합니다. 예: AI 반도체 투자 흐름")
    add_step(doc, "수동 대본은 비워둡니다.")
    add_step(doc, "스토리보드 생성 버튼을 누릅니다.")
    add_step(doc, "오른쪽 진행 상태 로그에서 생성 완료 메시지를 확인합니다.")
    add_step(doc, "아래 스토리보드 영역에 씬들이 생겼는지 확인합니다.")
    add_step(doc, "각 씬의 나레이션을 읽고 어색한 문장을 수정합니다.")
    add_step(doc, "각 씬의 이미지 프롬프트를 확인하고 원하는 장면 묘사로 수정합니다.")
    add_note(doc, "처음에는 짧은 주제로 테스트하세요. 결과가 마음에 들면 같은 방식으로 긴 작업을 진행합니다.")

    add_h1(doc, "5. 이미 써둔 대본으로 스토리보드 만드는 절차")
    add_step(doc, "스토리보드 생성 탭을 엽니다.")
    add_step(doc, "작업 이름을 입력합니다.")
    add_step(doc, "수동 대본 칸에 직접 작성한 대본을 붙여넣습니다.")
    add_step(doc, "AI가 씬 구성을 다시 만들게 하려면 스토리보드 생성을 누릅니다.")
    add_step(doc, "문장 기준으로만 나누고 싶으면 대본만 씬 분할을 누릅니다.")
    add_step(doc, "생성된 씬의 나레이션과 이미지 프롬프트를 수정합니다.")
    add_step(doc, "프로젝트 저장을 눌러 중간 저장합니다.")

    add_h1(doc, "6. 유튜브 자막을 가져와 영상 대본으로 쓰는 절차")
    add_step(doc, "왼쪽 메뉴에서 자막 추출 탭을 누릅니다.")
    add_step(doc, "YouTube URL 칸에 영상 주소를 붙여넣습니다.")
    add_step(doc, "언어 코드는 기본값 ko,en을 그대로 둡니다.")
    add_step(doc, "자막 추출하기 버튼을 누릅니다.")
    add_step(doc, "추출된 자막이 아래 텍스트 영역에 표시되는지 확인합니다.")
    add_step(doc, "필요하면 클립보드 복사 버튼으로 따로 저장합니다.")
    add_step(doc, "스토리보드로 보내기 버튼을 누릅니다.")
    add_step(doc, "스토리보드 생성 탭으로 이동한 뒤 수동 대본 칸에 자막이 들어갔는지 확인합니다.")
    add_step(doc, "스토리보드 생성 또는 대본만 씬 분할을 눌러 작업을 이어갑니다.")
    add_note(doc, "영상에 공개 자막이 없거나 자동 자막 접근이 막혀 있으면 추출이 실패할 수 있습니다.")

    add_h1(doc, "7. 씬 이미지 생성 절차")
    add_step(doc, "스토리보드 생성 탭에서 씬 목록이 있는지 확인합니다.")
    add_step(doc, "왼쪽 API 설정에서 Gemini API Key가 저장되어 있는지 확인합니다.")
    add_step(doc, "이미지 화풍을 원하는 스타일로 선택하고 저장합니다.")
    add_step(doc, "전체 이미지 생성 버튼을 누릅니다.")
    add_step(doc, "진행 로그에서 씬 1, 씬 2처럼 순서대로 생성되는지 확인합니다.")
    add_step(doc, "각 씬 왼쪽에 이미지 미리보기가 표시되는지 확인합니다.")
    add_step(doc, "마음에 들지 않는 씬은 해당 씬의 이미지 버튼을 눌러 다시 생성합니다.")
    add_note(doc, "좋은 이미지를 얻으려면 이미지 프롬프트에 인물, 배경, 분위기, 구도를 구체적으로 적습니다.")

    add_h1(doc, "8. 영상 파일로 내보내는 절차")
    add_step(doc, "스토리보드 생성 탭에서 이미지가 생성된 씬이 있는지 확인합니다.")
    add_step(doc, "MP4/WebM 내보내기 버튼을 누릅니다.")
    add_step(doc, "브라우저가 렌더링을 시작하면 진행 로그를 기다립니다.")
    add_step(doc, "완료되면 브라우저 다운로드 파일을 확인합니다.")
    add_step(doc, "MP4가 지원되지 않는 브라우저에서는 WebM으로 저장될 수 있습니다.")
    add_step(doc, "렌더링을 중단해야 하면 렌더 취소 요청 버튼을 누릅니다.")
    add_note(doc, "영상 내보내기는 브라우저 성능에 영향을 받습니다. 긴 영상은 시간이 걸릴 수 있습니다.")

    add_h1(doc, "9. 기존 영상 자동화 절차")
    add_step(doc, "왼쪽 메뉴에서 영상 자동화 탭을 누릅니다.")
    add_step(doc, "프로젝트 이름을 입력합니다.")
    add_step(doc, "영상 경로/URL 칸에 유튜브 주소나 로컬 영상 파일 경로를 입력합니다.")
    add_step(doc, "처음에는 고급 설정을 건드리지 않고 자동화 시작을 누릅니다.")
    add_step(doc, "자동화 로그에서 처리 상태를 확인합니다.")
    add_step(doc, "후보 장면이 표시되면 결과를 확인합니다.")
    add_step(doc, "최종 영상 합성 버튼을 눌러 자막 포함 결과 영상을 만듭니다.")
    add_h2(doc, "고급 설정 사용 기준")
    add_bullet(doc, "SALIENCY: 중요한 장면 위주로 찾고 싶을 때 사용합니다.")
    add_bullet(doc, "UNIFORM: 영상 전체에서 균일하게 장면을 뽑고 싶을 때 사용합니다.")
    add_bullet(doc, "결과 수: 후보 장면을 몇 개 표시할지 정합니다.")
    add_bullet(doc, "샘플링 간격: 영상을 몇 초마다 검사할지 정합니다.")
    add_bullet(doc, "최소 간격: 후보 장면끼리 너무 붙지 않도록 제한합니다.")

    add_h1(doc, "10. 프로젝트 저장과 불러오기")
    add_step(doc, "스토리보드 작업 중간이나 완료 후 프로젝트 저장 버튼을 누릅니다.")
    add_step(doc, "나중에 다시 열 때는 저장된 프로젝트 탭을 누릅니다.")
    add_step(doc, "원하는 프로젝트의 불러오기 버튼을 누릅니다.")
    add_step(doc, "스토리보드 생성 탭으로 이동하면 이전 씬이 복원됩니다.")
    add_step(doc, "필요 없는 프로젝트는 삭제 버튼으로 지웁니다.")
    add_note(doc, "저장 데이터는 브라우저 로컬 저장소에 있습니다. 브라우저 데이터를 삭제하면 같이 사라질 수 있습니다.")

    add_h1(doc, "11. 파일로 백업하는 방법")
    add_step(doc, "스토리보드 생성 탭에서 JSON 내보내기를 누릅니다.")
    add_step(doc, "다운로드된 storyboard.json 파일을 별도 폴더에 보관합니다.")
    add_step(doc, "표 형태로 보고 싶으면 CSV 내보내기를 누릅니다.")
    add_step(doc, "완성 영상은 MP4/WebM 내보내기로 저장합니다.")
    add_note(doc, "중요한 프로젝트는 프로젝트 저장만 믿지 말고 JSON 파일도 따로 보관하세요.")

    add_h1(doc, "12. 자주 막히는 문제")
    add_table(
        doc,
        [
            ("브라우저가 안 열림", "직접 http://127.0.0.1:3000/ 주소를 입력합니다."),
            ("스토리보드 생성 실패", "Gemini API Key 저장 여부와 테스트 결과를 확인합니다."),
            ("이미지 생성 실패", "이미지 모델, API Key, 프롬프트 길이를 확인합니다."),
            ("자막 추출 실패", "영상에 공개 자막이 있는지 확인하고 언어 코드를 바꿔봅니다."),
            ("MP4 대신 WebM 저장", "브라우저의 MP4 인코딩 지원 차이입니다. WebM 파일을 사용합니다."),
            ("저장 프로젝트가 사라짐", "브라우저 데이터 삭제 여부를 확인합니다."),
            ("영상 자동화 실패", "런처의 백엔드 서버가 실행 중인지 확인합니다."),
        ],
    )

    add_h1(doc, "13. 추천 작업 흐름")
    add_h2(doc, "유튜브 영상을 참고해 새 쇼츠 만들기")
    add_step(doc, "자막 추출 탭에서 유튜브 자막을 가져옵니다.")
    add_step(doc, "스토리보드로 보내기를 누릅니다.")
    add_step(doc, "스토리보드 생성 탭에서 씬을 만듭니다.")
    add_step(doc, "나레이션과 이미지 프롬프트를 정리합니다.")
    add_step(doc, "전체 이미지 생성을 누릅니다.")
    add_step(doc, "프로젝트 저장을 누릅니다.")
    add_step(doc, "MP4/WebM 내보내기로 결과 파일을 만듭니다.")

    add_h2(doc, "직접 쓴 대본으로 이미지 영상 만들기")
    add_step(doc, "스토리보드 생성 탭에서 작업 이름을 입력합니다.")
    add_step(doc, "수동 대본 칸에 대본을 붙여넣습니다.")
    add_step(doc, "대본만 씬 분할 또는 스토리보드 생성을 누릅니다.")
    add_step(doc, "각 씬 이미지 프롬프트를 구체적으로 고칩니다.")
    add_step(doc, "전체 이미지 생성을 누릅니다.")
    add_step(doc, "필요하면 JSON으로 백업합니다.")
    add_step(doc, "MP4/WebM으로 내보냅니다.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_h1(doc, "14. 화면별 핵심 버튼 요약")
    add_table(
        doc,
        [
            ("스토리보드 생성", "새 대본을 씬으로 만들고 이미지와 영상 내보내기를 진행합니다."),
            ("영상 자동화", "기존 영상에서 후보 장면과 자막 처리를 자동 실행합니다."),
            ("자막 추출", "유튜브 자막을 가져와 대본으로 보냅니다."),
            ("저장된 프로젝트", "저장한 작업을 불러오거나 삭제합니다."),
            ("저장", "API 설정 또는 프로젝트 내용을 브라우저에 저장합니다."),
            ("테스트", "Gemini API Key가 정상인지 확인합니다."),
            ("전체 이미지 생성", "모든 씬의 이미지를 한 번에 생성합니다."),
            ("MP4/WebM 내보내기", "브라우저에서 최종 영상 파일을 만듭니다."),
        ],
    )

    doc.save(OUT_PATH)
    TXT_PATH.write_text("\r\n".join(TEXT_LINES).strip() + "\r\n", encoding="utf-8-sig")
    print(OUT_PATH)
    print(TXT_PATH)


if __name__ == "__main__":
    build_manual()
